"""
Output Rendering & Report Generation Demo
--------------------------------------------
Run with:  streamlit run report_generation_app.py

Demonstrates:
- Rich in-app output rendering (text, tables, metrics, charts, code, JSON)
- Generating downloadable reports: PDF, Word (.docx), Excel (.xlsx), HTML, CSV
- Using st.download_button correctly (in-memory buffers, no temp files needed)
- Templated report content driven by user input
- Combining charts + tables + narrative text into a single exported report

Dependencies:
    pip install streamlit pandas numpy matplotlib fpdf2 python-docx openpyxl
"""

import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from io import BytesIO
from datetime import datetime

from fpdf import FPDF
from docx import Document
from docx.shared import Inches, Pt, RGBColor
import openpyxl
from openpyxl.styles import Font, PatternFill
from openpyxl.utils.dataframe import dataframe_to_rows

st.set_page_config(page_title="Output Rendering & Reports", page_icon="🧾", layout="wide")

st.title("🧾 Output Rendering & Report Generation")
st.caption("Render results in-app, then export them as PDF, Word, Excel, HTML, or CSV.")

# ============================================================================
# 1. Sample data + user-configurable report inputs
# ============================================================================
st.sidebar.header("Report Settings")
report_title = st.sidebar.text_input("Report title", "Quarterly Sales Report")
author_name = st.sidebar.text_input("Prepared by", "Analytics Team")
n_rows = st.sidebar.slider("Sample data rows", 10, 200, 60)


@st.cache_data
def generate_data(n):
    rng = np.random.default_rng(3)
    dates = pd.date_range("2024-01-01", periods=n, freq="D")
    df = pd.DataFrame(
        {
            "Date": dates,
            "Region": rng.choice(["North", "South", "East", "West"], size=n),
            "Sales": rng.normal(500, 80, size=n).round(2),
            "Units": rng.integers(20, 100, size=n),
        }
    )
    return df


df = generate_data(n_rows)
summary = df.groupby("Region", as_index=False).agg(
    Total_Sales=("Sales", "sum"), Total_Units=("Units", "sum"), Avg_Sale=("Sales", "mean")
).round(2)

# ============================================================================
# 2. In-app output rendering (many st output types)
# ============================================================================
st.header("1. In-App Output Rendering")

col1, col2, col3 = st.columns(3)
col1.metric("Total Sales", f"${df['Sales'].sum():,.0f}")
col2.metric("Total Units", f"{df['Units'].sum():,}")
col3.metric("Avg Sale", f"${df['Sales'].mean():,.2f}")

tab_table, tab_chart, tab_code, tab_json = st.tabs(["Table", "Chart", "Code", "JSON"])

with tab_table:
    st.dataframe(summary, use_container_width=True)

with tab_chart:
    fig, ax = plt.subplots(figsize=(8, 3.5))
    ax.bar(summary["Region"], summary["Total_Sales"], color="#6366f1")
    ax.set_ylabel("Total Sales")
    ax.set_title("Sales by Region")
    for spine in ["top", "right"]:
        ax.spines[spine].set_visible(False)
    st.pyplot(fig)

with tab_code:
    st.code(
        "summary = df.groupby('Region').agg(Total_Sales=('Sales','sum'))",
        language="python",
    )

with tab_json:
    st.json(summary.to_dict(orient="records"))

st.markdown("---")

# ============================================================================
# 3. Report generation functions
# ============================================================================
st.header("2. Generate & Download Reports")
st.write("Each export below is built in-memory (no temp files) and streamed via `st.download_button`.")


def make_chart_image(summary_df) -> BytesIO:
    """Render a matplotlib chart to a PNG buffer for embedding in reports."""
    fig, ax = plt.subplots(figsize=(6, 3))
    ax.bar(summary_df["Region"], summary_df["Total_Sales"], color="#4338ca")
    ax.set_title("Total Sales by Region")
    ax.set_ylabel("Sales")
    for spine in ["top", "right"]:
        ax.spines[spine].set_visible(False)
    buf = BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


def build_pdf(title, author, df, summary_df) -> bytes:
    """Build a PDF report using fpdf2."""
    chart_buf = make_chart_image(summary_df)
    chart_path = "/tmp/_report_chart.png"
    with open(chart_path, "wb") as f:
        f.write(chart_buf.getbuffer())

    pdf = FPDF()
    pdf.add_page()

    pdf.set_font("Helvetica", "B", 18)
    pdf.cell(0, 12, title, ln=True)
    pdf.set_font("Helvetica", "", 10)
    pdf.set_text_color(100, 100, 100)
    pdf.cell(0, 8, f"Prepared by {author} | Generated {datetime.now():%Y-%m-%d %H:%M}", ln=True)
    pdf.set_text_color(0, 0, 0)
    pdf.ln(4)

    pdf.set_font("Helvetica", "B", 13)
    pdf.cell(0, 10, "Summary by Region", ln=True)
    pdf.image(chart_path, w=170)
    pdf.ln(4)

    # Summary table
    pdf.set_font("Helvetica", "B", 10)
    col_widths = [45, 45, 45, 45]
    headers = list(summary_df.columns)
    for h, w in zip(headers, col_widths):
        pdf.cell(w, 8, str(h), border=1)
    pdf.ln()

    pdf.set_font("Helvetica", "", 10)
    for _, row in summary_df.iterrows():
        for val, w in zip(row, col_widths):
            pdf.cell(w, 8, str(val), border=1)
        pdf.ln()

    pdf.ln(6)
    pdf.set_font("Helvetica", "I", 9)
    pdf.multi_cell(0, 6, f"Report covers {len(df)} records from {df['Date'].min():%Y-%m-%d} to {df['Date'].max():%Y-%m-%d}.")

    return bytes(pdf.output(dest="S"))


def build_docx(title, author, df, summary_df) -> BytesIO:
    """Build a Word report using python-docx."""
    doc = Document()

    title_p = doc.add_heading(title, level=0)
    meta = doc.add_paragraph(f"Prepared by {author}  |  Generated {datetime.now():%Y-%m-%d %H:%M}")
    meta.runs[0].font.size = Pt(10)
    meta.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_heading("Summary by Region", level=1)

    chart_buf = make_chart_image(summary_df)
    doc.add_picture(chart_buf, width=Inches(6))

    table = doc.add_table(rows=1, cols=len(summary_df.columns))
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for i, col_name in enumerate(summary_df.columns):
        hdr_cells[i].text = str(col_name)

    for _, row in summary_df.iterrows():
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)

    doc.add_heading("Notes", level=1)
    doc.add_paragraph(
        f"This report covers {len(df)} records from {df['Date'].min():%Y-%m-%d} "
        f"to {df['Date'].max():%Y-%m-%d}."
    )

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def build_xlsx(df, summary_df) -> BytesIO:
    """Build an Excel workbook with raw data + styled summary sheet."""
    wb = openpyxl.Workbook()

    ws_summary = wb.active
    ws_summary.title = "Summary"
    header_fill = PatternFill(start_color="4338CA", end_color="4338CA", fill_type="solid")
    header_font = Font(color="FFFFFF", bold=True)

    for row in dataframe_to_rows(summary_df, index=False, header=True):
        ws_summary.append(row)
    for cell in ws_summary[1]:
        cell.fill = header_fill
        cell.font = header_font
    for col_cells in ws_summary.columns:
        max_len = max(len(str(c.value)) for c in col_cells)
        ws_summary.column_dimensions[col_cells[0].column_letter].width = max_len + 4

    ws_raw = wb.create_sheet("Raw Data")
    for row in dataframe_to_rows(df, index=False, header=True):
        ws_raw.append(row)
    for cell in ws_raw[1]:
        cell.fill = header_fill
        cell.font = header_font

    buf = BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf


def build_html(title, author, summary_df) -> str:
    """Build a simple styled HTML report."""
    table_html = summary_df.to_html(index=False, classes="report-table", border=0)
    html = f"""
    <html>
    <head>
    <style>
        body {{ font-family: Arial, sans-serif; margin: 40px; color: #1f2937; }}
        h1 {{ color: #4338ca; }}
        .meta {{ color: #6b7280; font-size: 0.9em; margin-bottom: 24px; }}
        .report-table {{ border-collapse: collapse; width: 100%; }}
        .report-table th {{ background: #4338ca; color: white; padding: 8px; text-align: left; }}
        .report-table td {{ padding: 8px; border-bottom: 1px solid #e5e7eb; }}
    </style>
    </head>
    <body>
        <h1>{title}</h1>
        <div class="meta">Prepared by {author} | Generated {datetime.now():%Y-%m-%d %H:%M}</div>
        <h2>Summary by Region</h2>
        {table_html}
    </body>
    </html>
    """
    return html


# ============================================================================
# 4. Export buttons
# ============================================================================
exp_col1, exp_col2, exp_col3, exp_col4, exp_col5 = st.columns(5)

with exp_col1:
    if st.button("🧾 Build PDF"):
        st.session_state["pdf_bytes"] = build_pdf(report_title, author_name, df, summary)
    if "pdf_bytes" in st.session_state:
        st.download_button(
            "Download PDF", data=st.session_state["pdf_bytes"],
            file_name="report.pdf", mime="application/pdf",
        )

with exp_col2:
    if st.button("📄 Build Word"):
        st.session_state["docx_buf"] = build_docx(report_title, author_name, df, summary)
    if "docx_buf" in st.session_state:
        st.download_button(
            "Download DOCX", data=st.session_state["docx_buf"],
            file_name="report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

with exp_col3:
    if st.button("📊 Build Excel"):
        st.session_state["xlsx_buf"] = build_xlsx(df, summary)
    if "xlsx_buf" in st.session_state:
        st.download_button(
            "Download XLSX", data=st.session_state["xlsx_buf"],
            file_name="report.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )

with exp_col4:
    html_report = build_html(report_title, author_name, summary)
    st.download_button(
        "🌐 Download HTML", data=html_report,
        file_name="report.html", mime="text/html",
    )

with exp_col5:
    csv_data = df.to_csv(index=False).encode("utf-8")
    st.download_button(
        "📑 Download CSV", data=csv_data,
        file_name="raw_data.csv", mime="text/csv",
    )

st.markdown("---")
st.caption(
    "Reports are built in-memory with fpdf2, python-docx, and openpyxl, then served via "
    "st.download_button — no server-side temp files required for download."
)