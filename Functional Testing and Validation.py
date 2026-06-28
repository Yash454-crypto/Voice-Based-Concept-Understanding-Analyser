"""
Functional Testing & Validation Suite
----------------------------------------
Run with:  pytest test_app.py -v

Demonstrates:
- Unit tests for pure business logic (data generation, validation, aggregation)
- Functional/UI tests using Streamlit's AppTest framework (simulates a real session)
- Input validation testing (valid/invalid/edge cases)
- Testing file-generating functions (PDF/DOCX/XLSX) for structural correctness
- Parametrized tests for covering multiple input scenarios
- Fixtures for shared setup

Dependencies:
    pip install pytest streamlit pandas numpy fpdf2 python-docx openpyxl
"""

import re
import io
import pytest
import pandas as pd
import numpy as np

from streamlit.testing.v1 import AppTest


# ============================================================================
# SECTION 1: Pure business logic (extracted so it's testable without Streamlit)
# ============================================================================

def generate_data(n: int, seed: int = 3) -> pd.DataFrame:
    rng = np.random.default_rng(seed)
    dates = pd.date_range("2024-01-01", periods=n, freq="D")
    return pd.DataFrame(
        {
            "Date": dates,
            "Region": rng.choice(["North", "South", "East", "West"], size=n),
            "Sales": rng.normal(500, 80, size=n).round(2),
            "Units": rng.integers(20, 100, size=n),
        }
    )


def summarize(df: pd.DataFrame) -> pd.DataFrame:
    return (
        df.groupby("Region", as_index=False)
        .agg(Total_Sales=("Sales", "sum"), Total_Units=("Units", "sum"), Avg_Sale=("Sales", "mean"))
        .round(2)
    )


EMAIL_PATTERN = r"^[\w\.-]+@[\w\.-]+\.\w+$"


def validate_profile(email: str, age: int, password: str) -> list[str]:
    """Returns a list of validation error messages (empty list = valid)."""
    errors = []
    if not re.match(EMAIL_PATTERN, email or ""):
        errors.append("Email format looks invalid.")
    if age is None or age <= 0:
        errors.append("Age must be greater than 0.")
    if age is not None and age > 130:
        errors.append("Age must be realistic (<= 130).")
    if not password or len(password) < 8:
        errors.append("Password must be at least 8 characters.")
    return errors


# ============================================================================
# SECTION 2: Fixtures
# ============================================================================

@pytest.fixture
def sample_df():
    return generate_data(n=50, seed=42)


@pytest.fixture
def sample_summary(sample_df):
    return summarize(sample_df)


# ============================================================================
# SECTION 3: Unit tests — data generation
# ============================================================================

class TestGenerateData:
    def test_returns_correct_row_count(self):
        df = generate_data(30)
        assert len(df) == 30

    def test_columns_present(self, sample_df):
        expected_cols = {"Date", "Region", "Sales", "Units"}
        assert expected_cols.issubset(sample_df.columns)

    def test_deterministic_with_same_seed(self):
        df1 = generate_data(20, seed=99)
        df2 = generate_data(20, seed=99)
        pd.testing.assert_frame_equal(df1, df2)

    def test_different_seed_gives_different_data(self):
        df1 = generate_data(20, seed=1)
        df2 = generate_data(20, seed=2)
        assert not df1["Sales"].equals(df2["Sales"])

    def test_units_within_expected_range(self, sample_df):
        assert sample_df["Units"].between(20, 100).all()

    @pytest.mark.parametrize("n", [0, 1, 5, 500])
    def test_handles_various_sizes(self, n):
        df = generate_data(n)
        assert len(df) == n


# ============================================================================
# SECTION 4: Unit tests — aggregation logic
# ============================================================================

class TestSummarize:
    def test_summary_has_one_row_per_region(self, sample_df, sample_summary):
        assert set(sample_summary["Region"]) == set(sample_df["Region"].unique())

    def test_total_sales_matches_sum(self, sample_df, sample_summary):
        expected_total = round(sample_df["Sales"].sum(), 2)
        actual_total = round(sample_summary["Total_Sales"].sum(), 2)
        assert actual_total == pytest.approx(expected_total, rel=1e-2)

    def test_avg_sale_is_nonnegative(self, sample_summary):
        assert (sample_summary["Avg_Sale"] >= 0).all()

    def test_empty_dataframe_raises_or_handles_gracefully(self):
        empty_df = pd.DataFrame(columns=["Date", "Region", "Sales", "Units"])
        result = summarize(empty_df)
        assert result.empty


# ============================================================================
# SECTION 5: Validation logic tests (valid / invalid / edge cases)
# ============================================================================

class TestValidateProfile:
    def test_valid_input_has_no_errors(self):
        errors = validate_profile("user@example.com", 30, "supersecret")
        assert errors == []

    @pytest.mark.parametrize(
        "email",
        ["not-an-email", "missing@domain", "@nodomain.com", "", "spaces in@email.com"],
    )
    def test_invalid_email_formats_rejected(self, email):
        errors = validate_profile(email, 30, "supersecret")
        assert "Email format looks invalid." in errors

    @pytest.mark.parametrize(
        "email",
        ["a@b.co", "first.last@example.com", "user+tag@example.io"],
    )
    def test_valid_email_formats_accepted(self, email):
        errors = validate_profile(email, 30, "supersecret")
        assert "Email format looks invalid." not in errors

    @pytest.mark.parametrize("age", [0, -1, -100])
    def test_non_positive_age_rejected(self, age):
        errors = validate_profile("user@example.com", age, "supersecret")
        assert "Age must be greater than 0." in errors

    def test_unrealistic_age_rejected(self):
        errors = validate_profile("user@example.com", 200, "supersecret")
        assert "Age must be realistic (<= 130)." in errors

    def test_boundary_age_accepted(self):
        # Edge cases right at the boundary should NOT trigger errors
        errors = validate_profile("user@example.com", 130, "supersecret")
        assert "Age must be realistic (<= 130)." not in errors

    @pytest.mark.parametrize("password", ["", "short", "1234567"])
    def test_short_password_rejected(self, password):
        errors = validate_profile("user@example.com", 30, password)
        assert any("Password" in e for e in errors)

    def test_multiple_errors_accumulate(self):
        errors = validate_profile("bad-email", -5, "123")
        assert len(errors) == 3


# ============================================================================
# SECTION 6: Functional / structural tests for report-generation outputs
# ============================================================================

class TestReportArtifacts:
    """
    These test the *structure* of generated files (size, headers, sheet names)
    rather than pixel-perfect rendering, since that's what's worth asserting on.
    """

    def test_xlsx_has_expected_sheets(self, sample_df, sample_summary):
        import openpyxl
        from openpyxl.utils.dataframe import dataframe_to_rows

        wb = openpyxl.Workbook()
        ws_summary = wb.active
        ws_summary.title = "Summary"
        for row in dataframe_to_rows(sample_summary, index=False, header=True):
            ws_summary.append(row)
        ws_raw = wb.create_sheet("Raw Data")
        for row in dataframe_to_rows(sample_df, index=False, header=True):
            ws_raw.append(row)

        buf = io.BytesIO()
        wb.save(buf)
        buf.seek(0)

        reloaded = openpyxl.load_workbook(buf)
        assert reloaded.sheetnames == ["Summary", "Raw Data"]
        assert reloaded["Summary"].max_row == len(sample_summary) + 1  # +1 header

    def test_docx_contains_table_and_heading(self, sample_summary):
        from docx import Document

        doc = Document()
        doc.add_heading("Test Report", level=0)
        table = doc.add_table(rows=1, cols=len(sample_summary.columns))
        for i, col in enumerate(sample_summary.columns):
            table.rows[0].cells[i].text = str(col)
        for _, row in sample_summary.iterrows():
            cells = table.add_row().cells
            for i, val in enumerate(row):
                cells[i].text = str(val)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        reloaded = Document(buf)
        assert reloaded.paragraphs[0].text == "Test Report"
        assert len(reloaded.tables) == 1
        assert len(reloaded.tables[0].rows) == len(sample_summary) + 1

    def test_pdf_output_is_nonempty_bytes(self, sample_summary):
        from fpdf import FPDF

        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Helvetica", "B", 14)
        pdf.cell(0, 10, "Test Report", ln=True)
        output = bytes(pdf.output(dest="S"))

        assert isinstance(output, bytes)
        assert len(output) > 100  # a real PDF has real header/structure bytes
        assert output.startswith(b"%PDF")  # valid PDF magic bytes


# ============================================================================
# SECTION 7: Functional UI tests using Streamlit's AppTest framework
# ============================================================================
# AppTest simulates a real Streamlit session: it runs the actual script,
# lets you "click" buttons / set widget values, and inspect rendered output.
#
# This requires the target script to exist on disk. Below we test a tiny
# inline app defined via AppTest.from_string, plus an example pattern for
# testing a real file (commented) for when app.py is in the same directory.

MINI_APP_SOURCE = """
import streamlit as st

st.title("Mini App")

name = st.text_input("Name", key="name")
age = st.number_input("Age", min_value=0, max_value=130, value=0, key="age")

if st.button("Submit"):
    if not name:
        st.error("Name is required")
    elif age <= 0:
        st.error("Age must be positive")
    else:
        st.success(f"Hello {name}, age {age}")

if "counter" not in st.session_state:
    st.session_state.counter = 0

if st.button("Increment"):
    st.session_state.counter += 1

st.write(f"Counter: {st.session_state.counter}")
"""


class TestAppFunctional:
    def test_app_runs_without_exceptions(self):
        at = AppTest.from_string(MINI_APP_SOURCE)
        at.run()
        assert not at.exception

    def test_title_renders(self):
        at = AppTest.from_string(MINI_APP_SOURCE)
        at.run()
        assert at.title[0].value == "Mini App"

    def test_submit_with_empty_name_shows_error(self):
        at = AppTest.from_string(MINI_APP_SOURCE)
        at.run()
        at.button[0].click().run()  # Submit clicked with default empty inputs
        assert len(at.error) == 1
        assert "Name is required" in at.error[0].value

    def test_submit_with_valid_input_shows_success(self):
        at = AppTest.from_string(MINI_APP_SOURCE)
        at.run()
        at.text_input(key="name").set_value("Alice").run()
        at.number_input(key="age").set_value(30).run()
        at.button[0].click().run()
        assert len(at.success) == 1
        assert "Hello Alice, age 30" in at.success[0].value

    def test_counter_increments_and_persists_across_reruns(self):
        at = AppTest.from_string(MINI_APP_SOURCE)
        at.run()
        increment_btn = [b for b in at.button if b.label == "Increment"][0]
        increment_btn.click().run()
        increment_btn = [b for b in at.button if b.label == "Increment"][0]
        increment_btn.click().run()
        assert at.session_state["counter"] == 2
        assert "Counter: 2" in at.markdown[-1].value or "Counter: 2" in str(at.get("text"))

    def test_zero_age_blocks_submission(self):
        at = AppTest.from_string(MINI_APP_SOURCE)
        at.run()
        at.text_input(key="name").set_value("Bob").run()
        at.button[0].click().run()
        assert any("Age must be positive" in e.value for e in at.error)


# Example pattern for testing a real app file on disk:
#
# def test_real_app_loads():
#     at = AppTest.from_file("session_state_app.py")
#     at.run()
#     assert not at.exception


if __name__ == "__main__":
    pytest.main([__file__, "-v"])